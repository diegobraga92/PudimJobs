"""PDF/DOCX CV import: extract raw text and structure it into a CVStructure.

Text extraction is deterministic (``pypdf`` for PDF, ``python-docx`` for DOCX).
Structuring follows ADR 011 and mirrors the ADR 006 philosophy: a rule-based
heuristic core always runs, and an optional LLM (OpenAI-compatible) is tried
first when configured. Any LLM failure degrades gracefully to the heuristics,
so CV import works out of the box without an API key.
"""

import json
import logging
import re
from io import BytesIO
from pathlib import Path

import httpx
from docx import Document
from pydantic import ValidationError
from pypdf import PdfReader

from app.schemas.master_cv import (
    CVStructure,
    EducationItem,
    ExperienceItem,
    ProjectItem,
)
from app.services.llm_config import LlmRuntimeConfig

logger = logging.getLogger(__name__)

# Cap the LLM prompt so a single structuring call stays cheap and fast.
_LLM_MAX_CHARS = 6_000

SUPPORTED_EXTENSIONS = {".pdf", ".docx"}

_MONTH = (
    r"(?:jan(?:uary)?|feb(?:ruary)?|mar(?:ch)?|apr(?:il)?|may|june?|july?|"
    r"aug(?:ust)?|sep(?:tember)?|oct(?:ober)?|nov(?:ember)?|dec(?:ember)?)"
)
_YEAR = r"(?:19|20)\d{2}"
_DATEMONTH = rf"(?:{_MONTH}\s*,?\s*{_YEAR}|\d{{1,2}}/\d{{2,4}}|{_YEAR})"

# A bare date range, e.g. "Jan 2020 - Present" or "2017 - 2021".
_DATE_RANGE_RE = re.compile(
    rf"({_DATEMONTH})\s*(?:-|\u2013|\u2014|to)\s*((?:present|now|current)|{_DATEMONTH})",
    re.IGNORECASE,
)
# Any single date-ish token (used to scrub dates out of header lines).
_DATE_RE = re.compile(rf"{_DATEMONTH}", re.IGNORECASE)
_YEAR_RE = re.compile(rf"{_YEAR}")
_URL_RE = re.compile(
    r"(?:https?://|www\.)[^\s]+|[A-Za-z0-9_.-]+\.(?:com|dev|io|org|net|github\.io)(?:/[^\s]*)?",
    re.IGNORECASE,
)

# Section headings are matched against a *whole* line (anchored fullmatch), so
# inline fragments like "Skills: Python, SQL" never steal a section.
_SECTION_HEADINGS: tuple[tuple[str, str], ...] = (
    (
        "summary",
        r"(?:professional\s+|career\s+)?summary|profile|about(?:\s+me)?|objective",
    ),
    (
        "experience",
        r"(?:work|professional|relevant|employment|career)\s+(?:history|experience)"
        r"|employment|experience",
    ),
    (
        "education",
        r"education|academics?|academic\s+background|qualifications?",
    ),
    (
        "skills",
        r"technical\s+skills|core\s+competenc\w+|skills\s*(?:&|and)?\s*"
        r"(?:tools|abilities)?|technologies|tech\s+stack|competenc\w+",
    ),
    (
        "projects",
        r"(?:personal\s+|side\s+)?projects|portfolio",
    ),
    (
        "ignore",
        r"languages?|certifications?|licenses?|interests?|hobbies?|references"
        r"|awards?|honors?|volunteering|publications?|courses?|extracurricular",
    ),
)

# Words that usually are not skills (heading leftovers, common CV filler).
_NON_SKILL_WORDS = {
    "a", "an", "the", "and", "or", "of", "in", "on", "at", "to", "for", "with",
    "skills", "skill", "technical", "technologies", "tools", "core",
    "competencies", "competency", "competences", "summary", "profile",
    "experience", "education", "projects", "certifications", "languages",
}

_MONTH_INDEX = {
    "jan": 1, "feb": 2, "mar": 3, "apr": 4, "may": 5, "jun": 6,
    "jul": 7, "aug": 8, "sep": 9, "oct": 10, "nov": 11, "dec": 12,
}

# Common degree markers, used to disambiguate "Degree, Institution" ordering.
_DEGREE_RE = re.compile(
    r"\b(?:b\.?s\.?c\.?|b\.?a\.?|m\.?s\.?c\.?|m\.?a\.?|m\.?b\.?a\.?|bachelor|master|"
    r"ph\.?\s*d|doctorate|mba|diploma|licenciat\w+|bacharel|mestrad\w+|doutorad\w+|"
    r"engenhari\w+|tecnic\w+|technician|degree|hnd|foundation|associate)\b",
    re.IGNORECASE,
)

_BULLET_RE = re.compile(
    r"^\s*(?:[\u2022\u25cf\u25e6\u25aa\u2023\u2043\u00b7*+\>\-\u2013\u2014]"
    r"|\d{1,2}[.)]|\u25aa)\s+(.+)$"
)
_SEPARATORS = (" | ", " @ ", " \u2014 ", " \u2013 ", " at ", " / ", " - ")


class CVParsingError(Exception):
    """Base class for CV parsing failures (mapped to HTTP 400)."""


class UnsupportedFileTypeError(CVParsingError):
    """The uploaded file's extension is not .pdf or .docx."""


class InvalidFileError(CVParsingError):
    """The file could not be read as a PDF/DOCX."""

# --- Text extraction -------------------------------------------------------


def extract_text_from_pdf(content: bytes) -> str:
    """Extract text from a PDF byte stream (one line per page text block)."""
    try:
        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
    except Exception as exc:
        raise InvalidFileError("Could not read PDF file") from exc
    return "\n".join(pages)


def extract_text_from_docx(content: bytes) -> str:
    """Extract text from a DOCX byte stream (paragraphs + table cells)."""
    try:
        document = Document(BytesIO(content))
    except Exception as exc:
        raise InvalidFileError("Could not read DOCX file") from exc
    lines = [p.text.strip() for p in document.paragraphs if p.text and p.text.strip()]
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                lines.append(" | ".join(cells))
    return "\n".join(lines)


def extract_text(filename: str, content: bytes) -> str:
    """Dispatch text extraction based on the file's extension."""
    suffix = Path(filename).suffix.lower()
    if suffix == ".pdf":
        return extract_text_from_pdf(content)
    if suffix == ".docx":
        return extract_text_from_docx(content)
    raise UnsupportedFileTypeError(
        f"Unsupported CV file type: {suffix or '(none)'}; use .pdf or .docx"
    )


# --- Rule-based structuring ------------------------------------------------


def _match_heading(line: str) -> str | None:
    cleaned = line.strip().strip(":#* \t|-\u2013\u2014").lower()
    if not cleaned:
        return None
    for kind, pattern in _SECTION_HEADINGS:
        if re.fullmatch(pattern, cleaned):
            return kind
    return None


def _bullet_text(line: str) -> str | None:
    match = _BULLET_RE.match(line)
    return match.group(1).strip() if match else None


def _scrub(text: str) -> str:
    """Remove URLs/dates from a header-like line and collapse whitespace."""
    cleaned = _URL_RE.sub(" ", text)
    cleaned = _DATE_RANGE_RE.sub(" ", cleaned)
    cleaned = _DATE_RE.sub(" ", cleaned)
    return re.sub(r"\s+", " ", cleaned).strip(" \t|,;:\u00b7\u2013\u2014")


def _normalize_date(value: str) -> str:
    value = value.strip()
    low = value.lower()
    if low in {"present", "now", "current"}:
        return "Present"
    match = re.fullmatch(
        rf"({_MONTH})\s*,?\s*({_YEAR})",
        value,
        re.IGNORECASE,
    )
    if match:
        month = _MONTH_INDEX.get(match.group(1).lower()[:3], 1)
        return f"{match.group(2)}-{month:02d}"
    if re.fullmatch(rf"{_YEAR}", value):
        return value
    match = re.fullmatch(r"(\d{1,2})/(\d{4})", value)
    if match:
        return f"{match.group(2)}-{int(match.group(1)):02d}"
    return value


def _extract_date_range(header: str) -> tuple[str | None, str | None]:
    match = _DATE_RANGE_RE.search(header)
    if match:
        return _normalize_date(match.group(1)), _normalize_date(match.group(2))
    single = _DATE_RE.search(header)
    if single:
        return _normalize_date(single.group(0)), None
    return None, None


def _split_title_company(cleaned: str) -> tuple[str, str]:
    for sep in _SEPARATORS:
        if sep in cleaned:
            parts = [p.strip(" \t|,;") for p in cleaned.split(sep)]
            parts = [p for p in parts if p]
            if len(parts) >= 2:
                return parts[0][:255], parts[-1][:255]
    return cleaned[:255], ""


def _is_date_only(line: str) -> bool:
    """True when a line carries only a date/date-range (no title/company)."""
    return not _scrub(line)


def _merge_date_range(item: ExperienceItem, line: str) -> None:
    """Back-fill a standalone date line onto an already-flushed item."""
    start, end = _extract_date_range(line)
    if start and not item.start_date:
        item.start_date = start
    if end and not item.end_date:
        item.end_date = end


def _parse_experience(lines: list[str]) -> list[ExperienceItem]:
    items: list[ExperienceItem] = []
    header_parts: list[str] = []
    bullets: list[str] = []

    def flush() -> None:
        nonlocal header_parts, bullets
        if header_parts:
            items.append(_experience_from_header(" ".join(header_parts), bullets))
        header_parts, bullets = [], []

    for line in lines:
        bullet = _bullet_text(line)
        if bullet is not None:
            bullets.append(bullet)
            continue
        if _is_date_only(line):
            # A date on its own line belongs to the current header, or (when it
            # followed a flushed item) to that previous item — never a new one.
            if header_parts:
                header_parts.append(line)
            elif items:
                _merge_date_range(items[-1], line)
            continue
        flush()
        header_parts.append(line)
    flush()
    return items


def _experience_from_header(header: str, bullets: list[str]) -> ExperienceItem:
    start, end = _extract_date_range(header)
    title, company = _split_title_company(_scrub(header))
    if not company:
        company = title
    if not title:
        title = company
    return ExperienceItem(
        title=title[:255],
        company=company[:255],
        start_date=start,
        end_date=end,
        bullets=bullets[:50],
    )


def _parse_education(lines: list[str]) -> list[EducationItem]:
    items: list[EducationItem] = []
    for line in lines:
        if _bullet_text(line) is not None:
            continue
        year_match = _YEAR_RE.search(line)
        year = year_match.group(0) if year_match else None
        cleaned = re.sub(r"\(\s*\)", "", _scrub(line))
        parts = [
            p.strip()
            for p in re.split(r"\s{2,}|\s*[,\-\u2013\u2014|]\s*", cleaned)
            if p.strip()
        ]
        if not parts:
            continue
        if len(parts) >= 2 and _DEGREE_RE.search(parts[-1]) and not _DEGREE_RE.search(parts[0]):
            degree, institution = parts[-1], parts[0]
        elif len(parts) >= 2:
            degree, institution = parts[0], parts[-1]
        else:
            degree = institution = parts[0]
        items.append(
            EducationItem(
                institution=institution[:255],
                degree=degree[:255],
                year=year,
            )
        )
    return items


def _parse_skills(lines: list[str]) -> list[str]:
    seen: set[str] = set()
    skills: list[str] = []
    for line in lines:
        for token in re.split(r"[,;•|/\s]\s*|\s{2,}", line):
            token = token.strip(" \t.,;:()[]{}\u00b7\u2013\u2014")
            low = token.lower()
            if (
                not token
                or len(token) < 2
                or len(token) > 64
                or low in _NON_SKILL_WORDS
                or _YEAR_RE.fullmatch(token)
                or _URL_RE.fullmatch(token)
            ):
                continue
            if low not in seen:
                seen.add(low)
                skills.append(token)
        if len(skills) >= 200:
            break
    return skills


def _parse_projects(lines: list[str]) -> list[ProjectItem]:
    items: list[ProjectItem] = []
    name: str | None = None
    description_parts: list[str] = []
    pending_link: str | None = None

    def flush() -> None:
        nonlocal name, description_parts, pending_link
        if name is not None:
            description = " ".join(description_parts).strip()
            link_match = _URL_RE.search(description)
            link = link_match.group(0) if link_match else None
            name_clean = _URL_RE.sub(" ", name)
            if link is None and _URL_RE.search(name_clean):
                link = _URL_RE.search(name_clean).group(0)
            if link is None and pending_link:
                link = pending_link
            split = re.split(r"\s*(?:\u2014|\u2013|-)\s*", name_clean, maxsplit=1)
            if len(split) == 2 and split[1].strip():
                name_clean, inline_desc = split[0].strip(), split[1].strip()
                description = f"{inline_desc} {description}".strip()
            name_clean = re.sub(r"\s+", " ", name_clean).strip(" \t|,;:")
            if name_clean:
                items.append(
                    ProjectItem(
                        name=name_clean[:255],
                        description=description[:1000] or None,
                        link=link or None,
                    )
                )
        name, description_parts, pending_link = None, [], None

    for line in lines:
        if _URL_RE.fullmatch(line):
            # A URL on its own line links the project currently being built.
            pending_link = line
            continue
        bullet = _bullet_text(line)
        if bullet is not None:
            description_parts.append(bullet)
        else:
            flush()
            name = line.strip()
    flush()
    return items


def structure_cv(text: str) -> CVStructure:
    """Rule-based structuring of raw CV text into a CVStructure."""
    sections: dict[str, list[str]] = {
        "summary": [],
        "experience": [],
        "education": [],
        "skills": [],
        "projects": [],
    }
    current: str | None = None
    for line in (ln.strip() for ln in text.splitlines()):
        if not line:
            continue
        kind = _match_heading(line)
        if kind is not None:
            current = kind
            continue
        if current in sections:
            sections[current].append(line)

    return CVStructure(
        summary=" ".join(sections["summary"]).strip()[:1000],
        experience=_parse_experience(sections["experience"]),
        education=_parse_education(sections["education"]),
        skills=_parse_skills(sections["skills"]),
        projects=_parse_projects(sections["projects"]),
    )


# --- Optional LLM structuring (graceful degradation) -----------------------


def _strip_code_fence(content: str) -> str:
    stripped = content.strip()
    if stripped.startswith("```"):
        stripped = re.sub(r"^```(?:json)?\s*", "", stripped)
        stripped = re.sub(r"\s*```$", "", stripped)
    return stripped.strip()


async def _llm_structured(text: str, config: LlmRuntimeConfig) -> dict | None:
    """Call an OpenAI-compatible chat API; return the parsed JSON object or None."""
    schema = (
        '{"summary": "string", "experience": [{"company": "string", "title": "string", '
        '"start_date": "YYYY-MM or null", "end_date": "YYYY-MM or null", '
        '"bullets": ["string"]}], "education": [{"institution": "string", "degree": "string", '
        '"year": "string or null"}], "skills": ["string"], '
        '"projects": [{"name": "string", "description": "string or null", '
        '"link": "string or null"}]}'
    )
    system = (
        "You extract structured resume data from raw text. Return ONLY a JSON object "
        f"matching exactly this shape: {schema} "
        "Never invent facts, employers, dates, or skills that are not present in the "
        "text. Use null for missing dates and links."
    )
    user = f"Resume text:\n\n{text[:_LLM_MAX_CHARS]}"
    try:
        async with httpx.AsyncClient(timeout=30) as client:
            response = await client.post(
                f"{config.base_url.rstrip('/')}/chat/completions",
                headers={"Authorization": f"Bearer {config.api_key}"},
                json={
                    "model": config.model,
                    "messages": [
                        {"role": "system", "content": system},
                        {"role": "user", "content": user},
                    ],
                    "temperature": 0,
                    "max_tokens": 1600,
                },
            )
            response.raise_for_status()
            content = response.json()["choices"][0]["message"]["content"]
        parsed = json.loads(_strip_code_fence(content))
        return parsed if isinstance(parsed, dict) else None
    except Exception:
        logger.warning(
            "LLM CV structuring failed; falling back to rule-based parsing",
            exc_info=True,
        )
        return None


async def structure_cv_with_llm(text: str, config: LlmRuntimeConfig) -> CVStructure | None:
    """Structure ``text`` via the LLM, validated against CVStructure.

    Returns ``None`` when the API call fails or the reply is not a valid
    structure, so callers can fall back to the rule-based parser.
    """
    payload = await _llm_structured(text, config)
    if payload is None:
        return None
    try:
        return CVStructure.model_validate(payload)
    except ValidationError:
        logger.warning(
            "LLM CV structuring returned an invalid structure; using rule-based",
            exc_info=True,
        )
        return None


# --- Orchestration ---------------------------------------------------------


async def parse_cv_file(
    filename: str,
    content: bytes,
    *,
    llm_config: LlmRuntimeConfig | None = None,
) -> CVStructure:
    """Extract and structure a CV file into a CVStructure (never persists).

    Tries the LLM first when a config with an API key is supplied; any failure
    falls back to the deterministic rule-based parser.
    """
    text = extract_text(filename, content)
    if llm_config is not None and llm_config.enabled and llm_config.api_key:
        llm_result = await structure_cv_with_llm(text, llm_config)
        if llm_result is not None:
            return llm_result
    return structure_cv(text)

